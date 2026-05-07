"""
fix_manuscript.py
=================
Applies all pre-submission fixes to the manuscript .docx:
  4A. Replaces all 16 figure images with newly generated figures
  4B. Applies 14 targeted text edits (abstract, SHAP, p-values, captions, etc.)
  5.  Adds Sharpe ratio column to Table 10

Input:  C:/Users/Asus/Downloads/when_the_gate_stays_closed_FINAL_SUBMISSION (1).docx
Output: paper/when_the_gate_stays_closed_FINAL_SUBMISSION.docx  (in repo)

Run from repo root: python scripts/fix_manuscript.py
"""

from __future__ import annotations

import sys, os, shutil, zipfile, tempfile, re
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

# ── Paths ─────────────────────────────────────────────────────────────────────
REPO_ROOT   = Path(".")
FIG_DIR     = REPO_ROOT / "figures"
PAPER_DIR   = REPO_ROOT / "paper"
PAPER_DIR.mkdir(exist_ok=True)

# Input docx (user-provided FINAL_SUBMISSION)
INPUT_DOCX  = Path("C:/Users/Asus/Downloads/when_the_gate_stays_closed_FINAL_SUBMISSION (1).docx")
# Backup path
BACKUP_DOCX = PAPER_DIR / "when_the_gate_stays_closed_BACKUP.docx"
# Intermediate (images replaced)
INTERIM_DOCX = PAPER_DIR / "when_the_gate_stays_closed_INTERIM.docx"
# Final output
OUTPUT_DOCX  = PAPER_DIR / "when_the_gate_stays_closed_FINAL_SUBMISSION.docx"

# ── Ground truth values (from regenerate_all_figures.py run) ─────────────────
MEAN_IC      = -0.0005106    # canonical IC from ic_comparison
IC_STD       = 0.2204
ICIR         = -0.002317
T_HAC        = -0.09
P_HAC        = 0.536         # upper tail
N_DAYS       = 1512
TOPK1_SHARPE = -0.1602
EW_SHARPE    = 0.9575
MOM_SHARPE   = 0.5717
DM_STAT      = 0.4234
DM_P         = 0.6721
PERM_SHARPE_P = 0.742

# IC-level permutation p-values (from regenerate_all_figures.py block bootstrap)
P_PERM_A     = 0.599
P_PERM_B     = 0.601

# SHAP values (from shap_mean_abs_by_fold.csv)
SHAP_TOP1_FEAT = "rolling 63-day volatility"
SHAP_TOP1_VAL  = 3.38   # ×10⁻³
SHAP_TOP2_FEAT = "EMA-50"
SHAP_TOP2_VAL  = 2.66   # ×10⁻³
SHAP_TOP3_FEAT = "OC body normalised"
SHAP_TOP3_VAL  = 2.54   # ×10⁻³
SHAP_RHO_MIN   = 0.13
SHAP_RHO_MAX   = 0.73
SHAP_RHO_MEAN  = 0.33

# N=100 values
N100_MEAN_IC = -0.00619
N100_T       = -1.6227
N100_P       = 0.052    # upper tail

# Momentum (from momentum_ic_gate_results.csv)
MOM_MEAN_IC  = 0.01044
MOM_HAC_T    = 1.348
MOM_P_UPPER  = 1.0 - 0.08879    # 1 - hac_p_value = 1 - 0.089 ≈ 0.911 for upper tail
# Actually: momentum gate results show hac_p_value = 0.0888 (one-sided, computed by script)
# This means p=0.089 is already the upper-tail p (H1: IC > 0)
MOM_P_UPPER  = 0.089

changes_log = []

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 4A — Replace all 16 figure images in the .docx
# ═══════════════════════════════════════════════════════════════════════════════

print("\n" + "="*65)
print("  STEP 4A: Replacing figure images in manuscript")
print("="*65)

# Backup
shutil.copy2(str(INPUT_DOCX), str(BACKUP_DOCX))
print(f"  Backup created: {BACKUP_DOCX}")

# Image slot → new figure file mapping (in document order)
image_map = {
    "image1.png":  "fig01_hac_bandwidth.png",
    "image2.png":  "fig02_power_analysis.png",
    "image3.png":  "fig03_fold_level_ic.png",
    "image4.png":  "fig04_strategy_performance.png",
    "image5.png":  "fig05_sharpe_vs_k.png",
    "image6.png":  "fig06_permutation_ic.png",
    "image7.png":  "fig07_permutation_sharpe.png",
    "image8.png":  "fig08_subperiod_heatmap.png",
    "image9.png":  "fig09_tc_sensitivity.png",
    "image10.png": "fig10_factor_regression.png",
    "image11.png": "fig11_universe_robustness.png",
    "image12.png": "fig12_shap_importance.png",
    "image13.png": "fig13_dm_test.png",
    "image14.png": "fig14_vix_conditioned.png",
    "image15.png": "fig15_ic_gate_summary.png",
    "image16.png": "figA1_pipeline.png",
}

replaced_count = 0
with tempfile.TemporaryDirectory() as tmpdir:
    # Extract docx
    with zipfile.ZipFile(str(INPUT_DOCX), "r") as z:
        z.extractall(tmpdir)

    media_dir = Path(tmpdir) / "word" / "media"
    for old_img, new_fig in image_map.items():
        src = FIG_DIR / new_fig
        dst = media_dir / old_img
        if src.exists():
            shutil.copy2(str(src), str(dst))
            print(f"  ✓ {old_img} ← {new_fig} ({src.stat().st_size//1024} KB)")
            replaced_count += 1
        else:
            print(f"  ✗ MISSING: {src}")

    # Repack to interim docx
    with zipfile.ZipFile(str(INTERIM_DOCX), "w", zipfile.ZIP_DEFLATED) as zout:
        for root, dirs, files in os.walk(tmpdir):
            for fname in files:
                filepath = Path(root) / fname
                arcname  = filepath.relative_to(tmpdir)
                zout.write(str(filepath), str(arcname))

print(f"\n  Replaced {replaced_count}/16 images → {INTERIM_DOCX}")
changes_log.append(f"4A: {replaced_count}/16 figures replaced in docx")

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 4B — Text edits
# ═══════════════════════════════════════════════════════════════════════════════

print("\n" + "="*65)
print("  STEP 4B: Applying text edits to manuscript")
print("="*65)

from docx import Document
from docx.oxml.ns import qn
import copy

doc = Document(str(INTERIM_DOCX))

def replace_text_in_doc(doc, old_text, new_text, label=""):
    """
    Replace old_text with new_text across all paragraphs and table cells.
    Works at run level and at paragraph text level.
    Returns number of replacements.
    """
    count = 0

    def replace_in_para(para):
        nonlocal count
        if old_text not in para.text:
            return
        # Try run-level first
        for run in para.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                count += 1
        # If run-level didn't work (text split across runs), rebuild the paragraph
        if old_text in para.text:
            full = para.text
            if old_text in full:
                # Consolidate all runs into first run, clear rest
                new_full = full.replace(old_text, new_text)
                if para.runs:
                    para.runs[0].text = new_full
                    for run in para.runs[1:]:
                        run.text = ""
                    count += 1

    for para in doc.paragraphs:
        replace_in_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)

    if label:
        print(f"  {'✓' if count else '○'} {label}: {count} replacements")
    return count

def count_words(text):
    return len(text.split())

# ─────────────────────────────────────────────────────────────────────────────
# EDIT 1: Keywords — reduce to 6 keywords
# ─────────────────────────────────────────────────────────────────────────────
OLD_KW = ("IC-Gated Deployment Framework, false discovery risk reduction, machine learning, "
          "cross-sectional prediction, NASDAQ-100, walk-forward validation, ensemble learning, "
          "isotonic calibration, momentum positive control, ablation study, market efficiency")
NEW_KW = ("IC-Gated Deployment Framework, machine learning, cross-sectional prediction, "
          "walk-forward validation, NASDAQ-100, market efficiency")
n = replace_text_in_doc(doc, OLD_KW, NEW_KW, "Keywords reduced to 6")
changes_log.append(f"4B.1 Keywords: {n} replacements")

# ─────────────────────────────────────────────────────────────────────────────
# EDIT 2: mis-specification → misspecification
# ─────────────────────────────────────────────────────────────────────────────
n = replace_text_in_doc(doc, "mis-specification", "misspecification", "mis-specification → misspecification")
# Also hyphenated variants
n2 = replace_text_in_doc(doc, "mis‐specification", "misspecification")
changes_log.append(f"4B.2 Typo fix: {n+n2} replacements")

# ─────────────────────────────────────────────────────────────────────────────
# EDIT 3: Section 6.2 — SHAP top feature text (MFI-14 → rolling_vol_63d)
# ─────────────────────────────────────────────────────────────────────────────
OLD_SHAP = ("MFI-14 tops the list at 3.9 × 10−3, followed by MACD histogram "
            "(3.2 × 10−3) and rolling 63-day volatility (2.7 × 10−3), "
            "all within a narrow band consistent with noise-level attribution.")
NEW_SHAP = (f"{SHAP_TOP1_FEAT} tops the list at {SHAP_TOP1_VAL:.2f} × 10−3, "
            f"followed by {SHAP_TOP2_FEAT} ({SHAP_TOP2_VAL:.2f} × 10−3) and "
            f"{SHAP_TOP3_FEAT} ({SHAP_TOP3_VAL:.2f} × 10−3), all within a narrow "
            f"band consistent with noise-level attribution.")
n = replace_text_in_doc(doc, OLD_SHAP, NEW_SHAP, "SHAP top feature text updated")

# Simpler version if full string not found
if n == 0:
    n = replace_text_in_doc(doc, "MFI-14 tops the list at 3.9",
                            f"{SHAP_TOP1_FEAT} tops the list at {SHAP_TOP1_VAL:.2f}",
                            "SHAP top feature (short form)")
changes_log.append(f"4B.3 SHAP text: {n} replacements")

# ─────────────────────────────────────────────────────────────────────────────
# EDIT 4: SHAP ρ range text (mean ρ = 0.33 already corrected; add range)
# ─────────────────────────────────────────────────────────────────────────────
OLD_RHO = ("Mean Spearman rank correlation of the top-feature orderings across folds 9–12 "
           "is ρ = 0.33 — moderate rather than high")
NEW_RHO = (f"Spearman rank correlation of the top-feature orderings across folds 9–12 "
           f"ranges from ρ = {SHAP_RHO_MIN:.2f} to {SHAP_RHO_MAX:.2f} "
           f"(mean ρ = {SHAP_RHO_MEAN:.2f}) — moderate rather than high")
n = replace_text_in_doc(doc, OLD_RHO, NEW_RHO, "SHAP rho range added")

# Simpler fallback
if n == 0:
    n = replace_text_in_doc(doc, "is ρ = 0.33", f"ranges from ρ = {SHAP_RHO_MIN:.2f} to {SHAP_RHO_MAX:.2f} (mean ρ = {SHAP_RHO_MEAN:.2f})",
                            "SHAP rho (simple fallback)")
changes_log.append(f"4B.4 SHAP rho range: {n} replacements")

# ─────────────────────────────────────────────────────────────────────────────
# EDIT 5: Algorithm 1 Step 6 — add redundancy note
# ─────────────────────────────────────────────────────────────────────────────
OLD_ALG = "Gate condition A: t_HAC > 1.645 AND IC̅ > 0."
NEW_ALG = ("Gate condition A: t_HAC > 1.645 AND IC̅ > 0 "
           "(the second clause is redundant by construction of the one-tailed upper test; "
           "stated for implementation clarity).")
n = replace_text_in_doc(doc, OLD_ALG, NEW_ALG, "Algorithm 1 Step 6 redundancy note")
if n == 0:
    n = replace_text_in_doc(doc, "Gate condition A: t_HAC > 1.645",
                            "Gate condition A: t_HAC > 1.645 (see note below re redundancy of IC̄ > 0)",
                            "Algorithm 1 (short fallback)")
changes_log.append(f"4B.5 Algorithm 1: {n} replacements")

# ─────────────────────────────────────────────────────────────────────────────
# EDIT 6: Section 5.3 — Update IC-level permutation p-values
# (original said p=0.797 Type A, p=0.754 Type B; new values from block bootstrap)
# ─────────────────────────────────────────────────────────────────────────────
# Update Type A p-value
n = replace_text_in_doc(doc, "Type A (temporal permutation, left): p = 0.797",
                        f"Type A (IID bootstrap, left): p = {P_PERM_A:.3f}",
                        "IC perm Type A p-value in figure caption")
n2 = replace_text_in_doc(doc, "p = 0.797",
                         f"p = {P_PERM_A:.3f}",
                         "IC perm Type A p-value (generic)")
changes_log.append(f"4B.6a Perm p_typeA: {n+n2} replacements")

# Update Type B p-value
n = replace_text_in_doc(doc, "Type B (block permutation, block = 5 days, right): p = 0.754",
                        f"Type B (block bootstrap, block = 5 days, right): p = {P_PERM_B:.3f}",
                        "IC perm Type B p-value in figure caption")
n2 = replace_text_in_doc(doc, "p = 0.754",
                         f"p = {P_PERM_B:.3f}",
                         "IC perm Type B p-value (generic)")
changes_log.append(f"4B.6b Perm p_typeB: {n+n2} replacements")

# Update section 5.3 consolidation sentence
OLD_PERM_SECT = ("The Sharpe-based permutation p-value of 0.742 (Table 4) is a distinct test "
                 "from the IC-level Type A (0.797) and Type B (0.754) results above")
NEW_PERM_SECT = (f"The Sharpe-based permutation p-value of {PERM_SHARPE_P:.3f} (Table 4) is a "
                 f"distinct test from the IC-level Type A ({P_PERM_A:.3f}) and Type B "
                 f"({P_PERM_B:.3f}) results above")
n = replace_text_in_doc(doc, OLD_PERM_SECT, NEW_PERM_SECT, "Section 5.3 perm consolidation")
if n == 0:
    n = replace_text_in_doc(doc, "IC-level Type A (0.797) and Type B (0.754)",
                            f"IC-level Type A ({P_PERM_A:.3f}) and Type B ({P_PERM_B:.3f})",
                            "Section 5.3 type A/B inline (fallback)")
changes_log.append(f"4B.6c Perm consolidation sentence: {n} replacements")

# Add clarifying sentence after the consolidation
OLD_PERM_END = ("all three evaluate the null from different angles and are mutually consistent.")
NEW_PERM_END = ("all three evaluate the null from different angles and are mutually consistent. "
                f"For clarity: the IC-level Type A (p = {P_PERM_A:.3f}) and Type B "
                f"(p = {P_PERM_B:.3f}) tests form the gate's non-parametric criterion; "
                f"the Sharpe-based permutation (p = {PERM_SHARPE_P:.3f}, Table 4) is an "
                "independent diagnostic of strategy-level performance. All three confirm the "
                "gate-closed decision.")
n = replace_text_in_doc(doc, OLD_PERM_END, NEW_PERM_END, "Section 5.3 clarifying sentence")
changes_log.append(f"4B.6d Perm clarity sentence: {n} replacements")

# ─────────────────────────────────────────────────────────────────────────────
# EDIT 7: Table 4 caption — add footnote about permutation p
# ─────────────────────────────────────────────────────────────────────────────
OLD_T4 = ("Table 4. IC gate evaluation statistics across the full 1,512-day OOS window. "
          "ICIR = mean IC ÷ IC standard deviation.")
NEW_T4 = ("Table 4. IC gate evaluation statistics across the full 1,512-day OOS window. "
          "ICIR = mean IC ÷ IC standard deviation. "
          "† Permutation p reported is the Sharpe-based test (TopK1 Sharpe vs. null "
          "distribution); IC-level bootstrap p-values are "
          f"{P_PERM_A:.3f} (Type A) and {P_PERM_B:.3f} (Type B), "
          "both reported in Section 5.3 and Figure 6.")
n = replace_text_in_doc(doc, OLD_T4, NEW_T4, "Table 4 footnote added")
if n == 0:
    n = replace_text_in_doc(doc, "ICIR = mean IC ÷ IC standard deviation.",
                            "ICIR = mean IC ÷ IC standard deviation. "
                            f"†IC-level bootstrap p-values: Type A = {P_PERM_A:.3f}, Type B = {P_PERM_B:.3f} (Section 5.3).",
                            "Table 4 footnote (fallback)")
changes_log.append(f"4B.7 Table 4 footnote: {n} replacements")

# ─────────────────────────────────────────────────────────────────────────────
# EDIT 8: Section 5.2 — Period 2 interpretation with correct values
# ─────────────────────────────────────────────────────────────────────────────
OLD_P2 = ("The brief positive Sharpe in Period 2 reflects the general equity market’s "
          "strong recovery from the March 2020 trough — any long-only position would "
          "have benefited.")
NEW_P2 = ("The brief positive Sharpe in Period 2 (TopK1: +0.76) reflects broad market "
          "recovery from the March 2020 trough. TopK1’s Period 2 Sharpe substantially "
          "underperforms the equal-weight benchmark (+1.72) during the same period, "
          "confirming that concentration into a single ML-ranked stock captured less upside "
          "than equal weighting — consistent with an uninformative ranker.")
n = replace_text_in_doc(doc, OLD_P2, NEW_P2, "Section 5.2 Period 2 interpretation")
if n == 0:
    # Try shorter match
    n = replace_text_in_doc(doc,
        "any long-only position would have benefited.",
        "TopK1’s Period 2 Sharpe (+0.76) substantially underperforms "
        "the equal-weight benchmark (+1.72), consistent with an uninformative ranker.",
        "Period 2 (short fallback)")
changes_log.append(f"4B.8 Period 2 interpretation: {n} replacements")

# ─────────────────────────────────────────────────────────────────────────────
# EDIT 9: Contribution 2 — soften novelty claim
# ─────────────────────────────────────────────────────────────────────────────
OLD_C2 = ("ICGDF extends the purged walk-forward design of López de Prado (2018) with "
          "two elements not present in that framework: fold-specific isotonic calibration "
          "fitted on a dedicated held-out calibration window (never the test period), and the "
          "sequential [MLP val | Cal | Test] partitioning that prevents calibration from "
          "observing any test-period label distribution. These additions are necessary for "
          "the ECE validity claim.")
NEW_C2 = ("ICGDF extends the purged walk-forward design of López de Prado (2018) by "
          "formalising and integrating two elements into the ICGDF context: fold-specific "
          "isotonic calibration fitted on a dedicated held-out calibration window (never the "
          "test period), and the sequential [MLP val | Cal | Test] partitioning that prevents "
          "calibration from observing any test-period label distribution. This integration is "
          "necessary for the ECE validity claim.")
n = replace_text_in_doc(doc, OLD_C2, NEW_C2, "Contribution 2 novelty softened")
if n == 0:
    n = replace_text_in_doc(doc, "two elements not present in that framework:",
                            "two elements formalised within the ICGDF context:",
                            "Contribution 2 (short fallback)")
changes_log.append(f"4B.9 Contribution 2: {n} replacements")

# ─────────────────────────────────────────────────────────────────────────────
# EDIT 10: References — Harvey D. → Harvey David I.
# ─────────────────────────────────────────────────────────────────────────────
n = replace_text_in_doc(doc,
    "Harvey, D., Leybourne, S., & Newbold, P. (1997).",
    "Harvey, David I., Leybourne, S., & Newbold, P. (1997).",
    "Harvey D. → Harvey David I.")
if n == 0:
    n = replace_text_in_doc(doc, "Harvey, D., Leybourne",
                            "Harvey, David I., Leybourne",
                            "Harvey ref (short fallback)")
changes_log.append(f"4B.10 Harvey reference: {n} replacements")

# ─────────────────────────────────────────────────────────────────────────────
# EDIT 11: Figure 15 caption — clarify t-stat display
# ─────────────────────────────────────────────────────────────────────────────
OLD_F15 = ("Figure 15. IC gate summary panel. (Left) IC signal statistics across the full "
           "OOS window. (Center) Fold-level IC point estimates — no fold shows "
           "persistent directional bias; shaded regions/error bars represent block bootstrap "
           "95% confidence intervals (B = 2,000 resamples, block = 5 days), central markers "
           "represent fold-level point estimates. (Right) Gate decision across all robustness "
           "checks — the gate stays closed in all settings.")
NEW_F15 = ("Figure 15. IC gate summary panel. (Left) IC signal statistics across the full "
           "OOS window; the HAC t-statistic (−0.090) is shown at actual scale. "
           "(Center) Fold-level IC point estimates with 95% block bootstrap confidence "
           "intervals (B = 2,000 resamples, block = 5 days); all CIs span zero. "
           "(Right) IC gate decision across all robustness checks — gate stays closed "
           "in all settings.")
n = replace_text_in_doc(doc, OLD_F15, NEW_F15, "Figure 15 caption updated")
if n == 0:
    n = replace_text_in_doc(doc, "no fold shows persistent directional bias;",
                            "95% block bootstrap CIs shown; HAC t = −0.090 (actual scale, not ×10);",
                            "Fig 15 caption (short fallback)")
changes_log.append(f"4B.11 Figure 15 caption: {n} replacements")

# ─────────────────────────────────────────────────────────────────────────────
# EDIT 12: Figure 6 caption — update p-values and describe method
# ─────────────────────────────────────────────────────────────────────────────
OLD_F6 = ("Figure 6. Permutation null distributions for the IC-level test. "
          "Type A (temporal permutation, left): p = 0.797. "
          "Type B (block permutation, block = 5 days, right): p = 0.754. "
          "Both tests confirm the gate-closed decision. "
          "The observed mean IC lies at the center of the null distribution under "
          "both permutation schemes.")
NEW_F6 = (f"Figure 6. Bootstrap null distributions for the IC-level gate test (H₀: "
          f"μIC ≤ 0). "
          f"Type A (IID bootstrap, left): p = {P_PERM_A:.3f}. "
          f"Type B (block bootstrap, block = 5 days, right): p = {P_PERM_B:.3f}. "
          f"Both tests confirm the gate-closed decision. "
          f"The observed mean IC ({MEAN_IC:.5f}) is marked by the solid vertical line; "
          f"its position within the null distribution is consistent with the null "
          f"hypothesis of no predictive content.")
n = replace_text_in_doc(doc, OLD_F6, NEW_F6, "Figure 6 caption updated")
if n == 0:
    n = replace_text_in_doc(doc, "Type A (temporal permutation, left): p = 0.797",
                            f"Type A (IID bootstrap, left): p = {P_PERM_A:.3f}",
                            "Fig 6 caption Type A (short fallback)")
    n2 = replace_text_in_doc(doc, "Type B (block permutation, block = 5 days, right): p = 0.754",
                             f"Type B (block bootstrap, block = 5 days, right): p = {P_PERM_B:.3f}",
                             "Fig 6 caption Type B (short fallback)")
    n += n2
changes_log.append(f"4B.12 Figure 6 caption: {n} replacements")

# ─────────────────────────────────────────────────────────────────────────────
# EDIT 13: Figure 12 caption — add rho range
# ─────────────────────────────────────────────────────────────────────────────
OLD_F12 = ("Figure 12. Top-20 features by mean absolute tree-component SHAP value "
           "(CatBoost + Random Forest components, last 4 folds, TreeExplainer). "
           "Scores are uniformly small with no dominant feature. "
           "Mean inter-fold Spearman rank ρ = 0.33 (moderate) across tree-component "
           "SHAP attributions. "
           "Top features: MFI-14, MACD histogram, rolling 63-day volatility — all within "
           "a narrow noise-level band. MLP component excluded (requires kernel-based explainer).")
NEW_F12 = (f"Figure 12. Top-20 features by mean absolute tree-component SHAP value "
           f"(CatBoost + Random Forest components, last 4 folds, TreeExplainer). "
           f"Scores are uniformly small with no dominant feature. "
           f"Inter-fold Spearman rank ρ ranges from {SHAP_RHO_MIN:.2f} to "
           f"{SHAP_RHO_MAX:.2f} (mean {SHAP_RHO_MEAN:.2f}, moderate) across fold pairs. "
           f"Top features: {SHAP_TOP1_FEAT}, {SHAP_TOP2_FEAT}, {SHAP_TOP3_FEAT} — all "
           f"within a narrow noise-level band. MLP component excluded (requires kernel-based "
           f"explainer).")
n = replace_text_in_doc(doc, OLD_F12, NEW_F12, "Figure 12 caption updated")
if n == 0:
    n = replace_text_in_doc(doc, "Top features: MFI-14, MACD histogram",
                            f"Top features: {SHAP_TOP1_FEAT}, {SHAP_TOP2_FEAT}",
                            "Fig 12 caption features (short fallback)")
changes_log.append(f"4B.13 Figure 12 caption: {n} replacements")

# ─────────────────────────────────────────────────────────────────────────────
# EDIT 14: Fix any remaining p=0.464 → 0.536 (unified p-value)
# ─────────────────────────────────────────────────────────────────────────────
n = replace_text_in_doc(doc, "p = 0.464", f"p = {P_HAC:.3f}",
                        "p=0.464 → 0.536 (unified)")
n2 = replace_text_in_doc(doc, "p=0.464", f"p={P_HAC:.3f}", "p=0.464 → 0.536 (no space)")
changes_log.append(f"4B.14 p=0.464 unification: {n+n2} replacements")

# ─────────────────────────────────────────────────────────────────────────────
# Save the text-edited docx (interim, before Table 10 edit)
# ─────────────────────────────────────────────────────────────────────────────
PRE_T10_DOCX = PAPER_DIR / "when_the_gate_stays_closed_preT10.docx"
doc.save(str(PRE_T10_DOCX))
print(f"\n  Text edits saved to: {PRE_T10_DOCX}")

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 5 — Add Sharpe Ratio column to Table 10
# ═══════════════════════════════════════════════════════════════════════════════

print("\n" + "="*65)
print("  STEP 5: Adding Sharpe column to Table 10")
print("="*65)

from lxml import etree

doc2 = Document(str(PRE_T10_DOCX))

# Find Table 10 (contains "Momentum" and "(252d)" or "trailing")
table10 = None
for i, tbl in enumerate(doc2.tables):
    txt = " ".join(cell.text for row in tbl.rows for cell in row.cells)
    if ("Momentum" in txt or "momentum" in txt) and ("252" in txt or "trailing" in txt or "ML Ensemble" in txt or "ML" in txt):
        print(f"  Found candidate Table 10 at index {i}:")
        for row in tbl.rows:
            print(f"    {' | '.join(cell.text.strip()[:30] for cell in row.cells)}")
        table10 = tbl
        break

if table10 is None:
    print("  WARNING: Table 10 not found — searching all tables for 'Momentum'...")
    for i, tbl in enumerate(doc2.tables):
        txt = " ".join(cell.text for row in tbl.rows for cell in row.cells)
        if "Momentum" in txt:
            print(f"  Table {i} contains 'Momentum':")
            for row in tbl.rows:
                row_text = " | ".join(c.text.strip()[:25] for c in row.cells)
                print(f"    {row_text}")
else:
    # Add Sharpe column
    # Sharpe values to add per row
    sharpe_values = {
        "header":    "Sharpe (OOS)",
        "momentum":  f"{MOM_SHARPE:.2f}",   # 0.57
        "ml":        f"{TOPK1_SHARPE:.2f}",  # -0.16
    }

    def add_cell_to_row(row_elem, text, bold=False):
        """Add a new cell to an existing table row."""
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/"
        # Copy last cell's XML as template
        last_cell = row_elem.findall(qn("w:tc"))[-1]
        new_cell = copy.deepcopy(last_cell)
        # Clear text from new cell
        for p in new_cell.findall(qn("w:p")):
            for r in p.findall(qn("w:r")):
                for t in r.findall(qn("w:t")):
                    t.text = text
                if bold:
                    rpr = r.find(qn("w:rPr"))
                    if rpr is None:
                        rpr = etree.SubElement(r, qn("w:rPr"))
                    b_elem = rpr.find(qn("w:b"))
                    if b_elem is None:
                        etree.SubElement(rpr, qn("w:b"))
        row_elem.append(new_cell)

    rows = table10._tbl.findall(qn("w:tr"))
    for i, row_elem in enumerate(rows):
        cells = row_elem.findall(qn("w:tc"))
        # Get row text to identify which row this is
        row_text = " ".join(
            "".join(t.text or "" for t in c.iter(qn("w:t")))
            for c in cells
        ).lower()

        if i == 0 or "signal" in row_text or "mean ic" in row_text:
            add_cell_to_row(row_elem, sharpe_values["header"], bold=True)
            print(f"  Added header cell to row {i}: '{sharpe_values['header']}'")
        elif "momentum" in row_text:
            add_cell_to_row(row_elem, sharpe_values["momentum"])
            print(f"  Added Momentum Sharpe: {sharpe_values['momentum']}")
        elif "ml" in row_text or "ensemble" in row_text or "baseline" in row_text:
            add_cell_to_row(row_elem, sharpe_values["ml"])
            print(f"  Added ML Sharpe: {sharpe_values['ml']}")
        else:
            add_cell_to_row(row_elem, "—")
            print(f"  Added '—' to row {i}: {row_text[:40]}")

    changes_log.append("5: Sharpe column added to Table 10")

# ─────────────────────────────────────────────────────────────────────────────
# Save final output docx
# ─────────────────────────────────────────────────────────────────────────────
doc2.save(str(OUTPUT_DOCX))
print(f"\n  Final manuscript saved: {OUTPUT_DOCX}")
changes_log.append(f"OUTPUT: {OUTPUT_DOCX}")

# Clean up interim files
for f in [INTERIM_DOCX, PRE_T10_DOCX]:
    if f.exists():
        f.unlink()

# ═══════════════════════════════════════════════════════════════════════════════
# FINAL REPORT
# ═══════════════════════════════════════════════════════════════════════════════

print("\n" + "="*65)
print("  CHANGES APPLIED:")
print("="*65)
for c in changes_log:
    print(f"  ✓ {c}")

print("\n" + "="*65)
print("  VERIFICATION: Scanning final docx for key strings")
print("="*65)

doc_check = Document(str(OUTPUT_DOCX))
full_text = "\n".join(p.text for p in doc_check.paragraphs)
for table in doc_check.tables:
    for row in table.rows:
        for cell in row.cells:
            full_text += "\n" + " ".join(p.text for p in cell.paragraphs)

checks = [
    ("mean IC = −0.0005",           "−0.0005",     "Abstract IC value"),
    ("HAC t = −0.09",               "−0.09",        "HAC t-stat"),
    ("p = 0.536",                   "0.536",         "p-value"),
    ("Sharpe ratio of −0.16",       "−0.16",         "TopK1 Sharpe"),
    ("DM = 0.42",                   "0.42",          "DM statistic"),
    ("p = 0.672",                   "0.672",         "DM p-value"),
    ("misspecification",            "misspecification","Typo fixed"),
    ("rolling 63-day volatility",   "rolling 63-day", "SHAP top feature"),
    ("Harvey, David I.",            "Harvey, David I.","Reference fixed"),
    ("mis-specification REMOVED",   not "mis-specification" in full_text, "mis-specification gone"),
    (f"p = {P_PERM_A:.3f}",        f"{P_PERM_A:.3f}", "IC perm Type A p-value"),
]

all_pass = True
for label, search, desc in checks:
    if isinstance(search, bool):
        found = search
    else:
        found = search in full_text
    status = "✓" if found else "✗"
    if not found:
        all_pass = False
    print(f"  {status} {desc}: {'found' if found else 'NOT FOUND'}")

if all_pass:
    print("\n  ✓ All verification checks passed.")
else:
    print("\n  ✗ Some checks failed — review text edits.")

print(f"\n  Final docx: {OUTPUT_DOCX}")
print(f"  File size:  {OUTPUT_DOCX.stat().st_size/1024:.0f} KB")
